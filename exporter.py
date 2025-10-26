# exporter.py

import pandas as pd
from docx import Document


class Exporter:
    """
    Classe responsável por exportar dados e relatórios em diferentes formatos.
    """

    @staticmethod
    def exportar_para_csv(df: pd.DataFrame, caminho: str):
        try:
            df.to_csv(caminho, index=False, encoding="utf-8")
        except Exception as e:
            raise ValueError(f"Erro ao exportar para CSV: {e}")

    @staticmethod
    def exportar_para_json(df: pd.DataFrame, caminho: str):
        try:
            df.to_json(caminho, orient="records", force_ascii=False)
        except Exception as e:
            raise ValueError(f"Erro ao exportar para JSON: {e}")

    @staticmethod
    def exportar_para_word(df: pd.DataFrame, caminho: str):
        try:
            doc = Document()
            doc.add_heading("Relatório de Dados", level=1)

            tabela = doc.add_table(rows=1, cols=len(df.columns))
            tabela.style = "Table Grid"

            # Adiciona cabeçalhos
            for i, coluna in enumerate(df.columns):
                tabela.cell(0, i).text = coluna

            # Adiciona dados
            for _, linha in df.iterrows():
                row_cells = tabela.add_row().cells
                for i, valor in enumerate(linha):
                    row_cells[i].text = str(valor)

            doc.save(caminho)
        except Exception as e:
            raise ValueError(f"Erro ao exportar para Word: {e}")
