import streamlit as st
import pandas as pd
import plotly.express as px
from painel_inteligente import (
    gerar_alertas,
    gerar_metrica_cfop,
    gerar_metrica_sped,
    analise_inteligente,
    sugerir_correcao
)

import io
from docx import Document

st.set_page_config(page_title="Painel Inteligente Fiscal", layout="wide")
st.title("🦉 Painel Inteligente Fiscal")

# Upload único
arquivo = st.file_uploader("Envie um arquivo fiscal (CSV)", type="csv")

if arquivo:
    df = pd.read_csv(arquivo, sep=';')

    # Tabs organizadas
    aba = st.tabs([
        "🚨 Alertas",
        "📊 Métrica CFOP",
        "📁 Métrica SPED",
        "🧠 Insights Inteligentes",
        "🔧 Sugestões de Correção"
    ])

    # 🚨 Alertas
    with aba[0]:
        st.subheader("🚨 Alertas Fiscais")
        alertas = gerar_alertas(df)
        if alertas:
            for alerta in alertas:
                st.warning(alerta)

            # Exportar alertas em Word
            doc = Document()
            doc.add_heading("Alertas Fiscais", level=1)
            for alerta in alertas:
                doc.add_paragraph(alerta)

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button("📥 Baixar alertas em Word", buffer, file_name="alertas_fiscais.docx")
        else:
            st.success("✅ Nenhum alerta fiscal encontrado.")

    # 📊 CFOP
    with aba[1]:
        st.subheader("📊 CFOPs mais utilizados")
        metrica_cfop = gerar_metrica_cfop(df)
        if metrica_cfop.empty:
            st.warning("⚠️ Nenhuma coluna de CFOP encontrada.")
        else:
            st.dataframe(metrica_cfop)

            # Gráfico interativo
            fig = px.bar(metrica_cfop, x='CFOP', y='Quantidade', title='Distribuição de CFOPs')
            st.plotly_chart(fig)

    # 📁 SPED
    with aba[2]:
        st.subheader("📁 Vínculo com SPED Fiscal")
        metrica_sped = gerar_metrica_sped(df)
        if metrica_sped.empty:
            st.warning("⚠️ Nenhuma coluna relacionada ao SPED foi encontrada.")
        else:
            st.dataframe(metrica_sped)

    # 🧠 Insights
    with aba[3]:
        st.subheader("🧠 Análise Inteligente")
        insights = analise_inteligente(df)
        for item in insights:
            st.write(item)

    # 🔧 Sugestões de Correção
    with aba[4]:
        st.subheader("🔧 Sugestões de Correção Fiscal")
        sugestoes = sugerir_correcao(df)
        for item in sugestoes:
            st.write(item)