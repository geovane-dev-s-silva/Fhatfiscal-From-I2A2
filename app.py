# ChatFiscal - Vers√£o COM MEM√ìRIA INTELIGENTE + SANITIZA√á√ÉO

import os
import base64
import pandas as pd
import html
import re
from datetime import datetime
from dotenv import load_dotenv
import streamlit as st
from agent_manager import AgentManager
import logging
import sys
from gerar_pdf import gerar_relatorio_pdf
from visualizacao.interface import exibir_visualizacao

# CONFIGURA√á√ÉO INICIAL
def get_base64_icon(path):
    if os.path.exists(path):
        with open(path, "rb") as f:
            return base64.b64encode(f.read()).decode("utf-8")
    return None

# Substitua o emoji pela string data URI do mascote_corujito
icon_base64 = get_base64_icon("assets/mascote_corujito.png")
page_icon = f"data:image/png;base64,{icon_base64}" if icon_base64 else "ü¶â"

st.set_page_config(
    page_title="ChatFiscal",
    layout="wide",
    page_icon=page_icon
)
load_dotenv()
logger = logging.getLogger(__name__)

# Importa interface (com tratamento de erro)
sys.path.append(os.path.join(os.path.dirname(__file__), 'main'))
try:
    from main.interface import (
        montar_interface,
        exibir_resposta_agente,
        exibir_rodape,
        mostrar_alerta,
        mostrar_erro,
        mostrar_sucesso,
        boas_vindas,
        introducao_chatfiscal,
        exibir_dica_corujito
    )
    montar_interface()
except ImportError:
    pass

# Importa dicas inteligentes (com fallback)
try:
    from dicas_corujito import gerar_dica_corujito_inteligente
except ImportError:
    def gerar_dica_corujito_inteligente(df, contexto_empresa="", historico=None):
        return "An√°lise dos dados carregada com sucesso. Continue fazendo perguntas!"

# Inicializa o manager
manager = AgentManager()

# ESTILOS GLOBAIS
st.markdown("""
<style>
    .dica-corujito {
        background-color: #0D1B2A;
        border: 2px solid #FFD700;
        padding: 15px;
        border-radius: 10px;
        margin: 10px 0;
        box-shadow: 0 0 10px rgba(255, 215, 0, 0.3);
    }
    .dica-titulo {
        color: #FFD700;
        font-weight: bold;
        font-size: 16px;
        margin-bottom: 8px;
    }
    .dica-texto {
        color: #FFD700;
        font-size: 14px;
        font-family: monospace;
        line-height: 1.6;
    }
    .resposta-agente {
        background-color: #0D1B2A;
        border-left: 4px solid #FFD700;
        padding: 12px;
        border-radius: 5px;
        margin: 10px 0;
        color: #FFD700;
    }
</style>
""", unsafe_allow_html=True)

# SESSION STATE INITIALIZATION
if "dados_tabulares" not in st.session_state:
    st.session_state["dados_tabulares"] = []
if "past" not in st.session_state:
    st.session_state["past"] = []
if "generated" not in st.session_state:
    st.session_state["generated"] = []
if "pdf_carregado" not in st.session_state:
    st.session_state["pdf_carregado"] = False
if "arquivos_carregados" not in st.session_state:
    st.session_state["arquivos_carregados"] = set()
if "ultima_selecao" not in st.session_state:
    st.session_state["ultima_selecao"] = set()
if "pdf_list" not in st.session_state:
    st.session_state["pdf_list"] = []
if "df_csv_unificado" not in st.session_state:
    st.session_state["df_csv_unificado"] = None

# CARREGAMENTO DE AVATARES
def get_base64_image(path):
    try:
        if os.path.exists(path):
            with open(path, "rb") as img_file:
                return base64.b64encode(img_file.read()).decode("utf-8")
        else:
            logger.warning(f"Avatar n√£o encontrado: {path}")
            return None
    except Exception as e:
        logger.error(f"Erro ao carregar: {path} - {e}")
        return None

avatar_corujito_base64 = get_base64_image('assets/mascote_corujito.png')
avatar_agente_base64 = get_base64_image('assets/avatar_agente.png')
avatar_corujito_html = f"data:image/png;base64,{avatar_corujito_base64}" if avatar_corujito_base64 else None
avatar_agente_html = f"data:image/png;base64,{avatar_agente_base64}" if avatar_agente_base64 else None

def exibir_dica_popup(dica_texto):
    col1, col2 = st.columns([0.8, 10])
    with col1:
        if avatar_corujito_html:
            st.markdown(
                f'<img src="{avatar_corujito_html}" style="width:60px;height:60px;border-radius:50%;border:2px solid #FFD700;">',
                unsafe_allow_html=True
            )
        else:
            st.markdown("ü¶â")
    with col2:
        st.markdown(f"""
        <div class="dica-corujito">
            <div class="dica-titulo">üí° Dica do Corujito Fiscal:</div>
            <div class="dica-texto">{dica_texto}</div>
        </div>
        """, unsafe_allow_html=True)

# T√çTULO E ABAS
st.title("ChatFiscal - Seu Assistente Fiscal Inteligente")
abas = st.tabs(["Dados & Chat", "Hist√≥rico", "Auditoria", "Visualiza√ß√µes", "Painel"])

# ABA 0: DADOS & CHAT
with abas[0]:
    st.subheader("Upload de Arquivos")
    MAX_SIZE = 50 * 1024 * 1024
    arquivos = st.file_uploader(
        "Escolha seus arquivos fiscais (CSV, XML ou PDF)",
        type=["csv", "xml", "pdf"],
        accept_multiple_files=True
    )
    if arquivos:
        arquivos_names = {a.name for a in arquivos}
        if arquivos_names != st.session_state["ultima_selecao"]:
            st.session_state["ultima_selecao"] = arquivos_names
            st.session_state["dados_tabulares"] = []
            st.session_state["df_csv_unificado"] = None
            manager.arquivos_processados = set()
        arquivo_count = 0
        pdf_count = 0
        for arquivo in arquivos:
            nome = arquivo.name
            if arquivo.size > MAX_SIZE:
                st.error(f"Arquivo '{nome}' muito grande (m√°x 50MB)")
                continue
            if nome in manager.arquivos_processados:
                continue
            try:
                if nome.lower().endswith(".pdf"):
                    resultado = manager.carregar_arquivo(arquivo)
                    st.info(resultado if isinstance(resultado, str) else f"PDF '{nome}' processado")
                    pdf_count += 1
                    arquivo_count += 1
                    st.session_state["arquivos_carregados"].add(nome)
                elif nome.lower().endswith((".csv", ".xml")):
                    df = manager.carregar_arquivo(arquivo)
                    if isinstance(df, pd.DataFrame) and not df.empty:
                        arquivo_count += 1
                        st.session_state["arquivos_carregados"].add(nome)
            except Exception as e:
                st.error(f"Erro: {e}")
        if arquivo_count > 0:
            st.success(f"{arquivo_count} arquivo(s) carregado(s)")
        if pdf_count > 0:
            st.session_state["pdf_carregado"] = True

    df_unificado = None
    if st.session_state.get("dados_tabulares"):
        df_unificado = pd.concat(st.session_state["dados_tabulares"], ignore_index=True)
        df_unificado = df_unificado.drop_duplicates().reset_index(drop=True)
        st.session_state["df_csv_unificado"] = df_unificado
        with st.expander("Dados Carregados", expanded=True):
            st.dataframe(df_unificado, use_container_width=True)
            st.caption(f"Total: {len(df_unificado)} linhas | {len(df_unificado.columns)} colunas")
        st.markdown("---")
        try:
            with st.spinner("Corujito analisando seus dados..."):
                dica = gerar_dica_corujito_inteligente(df=df_unificado, contexto_empresa="SP | Com√©rcio")
                if dica:
                    exibir_dica_popup(dica)
        except Exception as e:
            logger.warning(f"Aviso na dica: {e}")
    else:
        st.info("Nenhum arquivo carregado ainda")

    st.markdown("---")
    st.subheader("Chat com o Agente Fiscal")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("Limpar conversa", use_container_width=True):
            st.session_state["past"] = []
            st.session_state["generated"] = []
            st.success("Conversa limpa!")
            st.rerun()
    with col2:
        if st.button("Limpar tudo", use_container_width=True):
            st.session_state.clear()
            manager.limpar_todos_pdfs()
            st.success("Tudo limpo!")
            st.rerun()

    st.markdown("---")
    if st.session_state["generated"]:
        st.write("Hist√≥rico da Conversa:")
        for i in range(len(st.session_state["generated"])):
            pergunta_sanitizada = html.escape(st.session_state['past'][i])
            resposta_sanitizada = html.escape(st.session_state['generated'][i])
            st.markdown(
                f"""
                <div style="display:flex;justify-content:flex-end;align-items:center;margin:10px 0;">
                    <div style="background-color:#1E3A5F;padding:12px 18px;border-radius:15px;
                                max-width:70%;color:white;box-shadow:0 2px 6px rgba(0,0,0,0.3);
                                font-size:15px;line-height:1.5;">
                        {pergunta_sanitizada}
                    </div>
                    <span style="font-size:42px;margin-left:10px;">ü§ì</span>
                </div>
                """, unsafe_allow_html=True,
            )
            if avatar_agente_html:
                st.markdown(
                    f"""
                    <div style="display:flex;align-items:flex-start;margin:10px 0;">
                        <img src="{avatar_agente_html}" 
                             style="width:48px;height:48px;border-radius:50%;
                                    margin-right:12px;border:2px solid #D4AF37;
                                    box-shadow:0 2px 4px rgba(0,0,0,0.2);">
                        <div style="background-color:#0D1B2A;border-left:4px solid #FFD700;
                                    padding:12px 16px;border-radius:8px;max-width:70%;
                                    color:#FFD700;box-shadow:0 2px 6px rgba(255,215,0,0.2);
                                    font-size:15px;line-height:1.6;white-space:pre-wrap;">
                            {resposta_sanitizada}
                        </div>
                    </div>
                    """, unsafe_allow_html=True,
                )
            else:
                st.markdown(
                    f"""
                    <div style="display:flex;align-items:flex-start;margin:10px 0;">
                        <span style="font-size:48px;margin-right:10px;">üë®‚Äçüíº</span>
                        <div style="background-color:#0D1B2A;border-left:4px solid #FFD700;
                                    padding:12px 16px;border-radius:8px;max-width:70%;
                                    color=#FFD700;font-size:15px;line-height:1.6;white-space:pre-wrap;">
                            {resposta_sanitizada}
                        </div>
                    </div>
                    """, unsafe_allow_html=True,
                )

    st.markdown("---")
    user_input = st.chat_input("Digite sua pergunta sobre os dados fiscais...")
    if user_input:
        user_input_limpo = re.sub(r'<[^>]+>', '', user_input).strip()
        if not user_input_limpo:
            st.warning("Pergunta n√£o pode estar vazia")
        elif not (df_unificado is not None or st.session_state.get("pdf_carregado")):
            st.error("Nenhum arquivo carregado! Fa√ßa upload primeiro.")
        else:
            with st.spinner("Analisando sua pergunta..."):
                try:
                    resposta = manager.gerar_resposta(user_input_limpo)
                    st.session_state["past"].append(user_input_limpo)
                    st.session_state["generated"].append(resposta)
                except Exception as e:
                    st.error(f"Erro ao processar: {e}")
                    logger.error(f"Erro: {e}")
            st.rerun()

# ABA 1: HIST√ìRICO
with abas[1]:
    st.subheader("Hist√≥rico da Sess√£o Atual")
    if st.session_state["past"] and st.session_state["generated"]:
        for i in range(len(st.session_state["past"])):
            with st.expander(f"{i+1}. {st.session_state['past'][i][:60]}{'...' if len(st.session_state['past'][i])>60 else ''}"):
                st.markdown(f"**Pergunta:** {st.session_state['past'][i]}")
                st.markdown(f"**Resposta:** {st.session_state['generated'][i]}")
                st.caption(f"{datetime.now().strftime('%d/%m/%Y %H:%M')}")
    else:
        st.info("Nenhuma conversa nesta sess√£o")

# ‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê
# ABA 2: AUDITORIA FISCAL
# ‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê
with abas[2]:
    st.subheader("üìã Auditoria Fiscal")
    
    # Fun√ß√£o auxiliar para detectar tipo de NF
    def detectar_tipo_nf(df):
        """Detecta se √© NFe, NFSe ou Misto"""
        if df is None or df.empty:
            return "N√£o especificado"
        
        tem_nfe = any(col.startswith(('emit_', 'dest_', 'ide_')) or 'cfop' in col.lower() for col in df.columns)
        tem_nfse = any(col.startswith(('prestador_', 'tomador_')) or 'iss' in col.lower() for col in df.columns)
        
        if tem_nfe and tem_nfse:
            return "MISTO (NFe + NFSe)"
        elif tem_nfe:
            return "NF-e (Nota Fiscal Eletr√¥nica)"
        elif tem_nfse:
            return "NFS-e (Nota Fiscal de Servi√ßo)"
        else:
            return "N√£o especificado"
    
    # M√©tricas principais
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        total_arquivos = len(st.session_state.get("arquivos_carregados", set()))
        st.metric("üìÅ Arquivos Processados", total_arquivos)
    
    with col2:
        total_registros = 0
        if st.session_state.get("dados_tabulares"):
            df_unificado = pd.concat(st.session_state["dados_tabulares"], ignore_index=True)
            total_registros = len(df_unificado)
        st.metric("üìä Registros Fiscais", total_registros)
    
    with col3:
        total_perguntas = len(st.session_state.get("past", []))
        st.metric("üí¨ Consultas Realizadas", total_perguntas)
    
    with col4:
        total_pdfs = len(st.session_state.get("pdf_list", []))
        st.metric("üìÑ Documentos PDF", total_pdfs)
    
    st.markdown("---")
    
    # Trilha de Auditoria - Hist√≥rico de A√ß√µes
    st.subheader("üîç Trilha de Auditoria")
    
    if st.session_state.get("historico"):
        with st.expander("üìú Hist√≥rico de Consultas", expanded=True):
            for i, (pergunta, resposta) in enumerate(st.session_state["historico"], 1):
                st.markdown(f"""
                **{i}. Consulta realizada:**
                - **Pergunta:** {pergunta}
                - **Data/Hora:** {datetime.now().strftime("%d/%m/%Y %H:%M")}
                - **Tipo:** An√°lise Fiscal Eletr√¥nica
                """)
                with st.expander(f"Ver resposta completa"):
                    st.write(resposta)
                st.markdown("---")
    else:
        st.info("üì≠ Nenhuma consulta realizada ainda")
    
    st.markdown("---")
    
    # Arquivos Processados
    st.subheader("üìÅ Arquivos Processados")
    
    if st.session_state.get("arquivos_carregados"):
        df_arquivos = pd.DataFrame({
            "Arquivo": list(st.session_state["arquivos_carregados"]),
            "Tipo": [nome.split(".")[-1].upper() for nome in st.session_state["arquivos_carregados"]],
            "Status": ["‚úÖ Processado"] * len(st.session_state["arquivos_carregados"])
        })
        st.dataframe(df_arquivos, use_container_width=True)
    else:
        st.info("üì≠ Nenhum arquivo carregado")
    
    st.markdown("---")
    
    # An√°lise de Qualidade dos Dados (COM EXPLICA√á√ÉO E TIPO DE NF)
    if st.session_state.get("dados_tabulares"):
        st.subheader("üìä Resumo da An√°lise de Qualidade")
        
        df_unificado = pd.concat(st.session_state["dados_tabulares"], ignore_index=True)
        
        # Detecta tipo de NF
        tipo_nf = detectar_tipo_nf(df_unificado)
        
        # Badge do tipo de NF
        if "MISTO" in tipo_nf:
            cor_badge = "üü°"
        elif "NF-e" in tipo_nf:
            cor_badge = "üîµ"
        elif "NFS-e" in tipo_nf:
            cor_badge = "üü£"
        else:
            cor_badge = "‚ö™"
        
        st.markdown(f"### {cor_badge} **Tipo de Documento:** {tipo_nf}")
        st.markdown("---")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.metric("Total de Campos", len(df_unificado.columns))
            st.metric("Total de Registros", len(df_unificado))
            st.metric("Registros Duplicados", df_unificado.duplicated().sum())
        
        with col2:
            # Calcula campos vazios
            total_celulas = len(df_unificado) * len(df_unificado.columns)
            campos_vazios = df_unificado.isnull().sum().sum()
            
            st.metric("Campos Vazios (Total)", f"{campos_vazios:,}")
            
            # Completude CORRETA
            completude = ((total_celulas - campos_vazios) / total_celulas * 100) if total_celulas > 0 else 0
            
            # Adiciona indicador de qualidade
            if completude >= 80:
                emoji = "üü¢"
                status = "Excelente"
            elif completude >= 60:
                emoji = "üü°"
                status = "Bom"
            elif completude >= 40:
                emoji = "üü†"
                status = "Moderado"
            else:
                emoji = "üî¥"
                status = "Cr√≠tico"
            
            st.metric("Completude dos Dados", f"{emoji} {completude:.1f}%", delta=status)
        
        # EXPLICA√á√ÉO CONTEXTUAL
        st.markdown("---")
        
        if "MISTO" in tipo_nf:
            st.info(f"""
            ‚ÑπÔ∏è **Explica√ß√£o - Dados MISTO (NFe + NFSe):**
            
            A completude de **{completude:.1f}%** √© **esperada** para dados consolidados de tipos diferentes:
            
            - **NF-e** possui campos espec√≠ficos (Emitente, Destinat√°rio, CFOP, ICMS, etc.)
            - **NFS-e** possui campos espec√≠ficos (Prestador, Tomador, ISS, etc.)
            - Quando consolidados, os campos exclusivos de cada tipo **ficam vazios naturalmente**
            
            **Exemplo pr√°tico:**
            - Registro NF-e: campos `prestador_*` ficam vazios (n√£o se aplicam)
            - Registro NFS-e: campos `emit_*`, `cfop`, `icms` ficam vazios (n√£o se aplicam)
            
            ‚úÖ **Isso N√ÉO indica problema de qualidade**, apenas estruturas diferentes!
            """)
        elif completude < 80:
            st.warning(f"""
            ‚ö†Ô∏è **Aten√ß√£o - Completude abaixo de 80%:**
            
            Campos vazios detectados: **{campos_vazios:,}** de {total_celulas:,} c√©lulas
            
            **Recomenda√ß√µes:**
            - Verifique se os XMLs est√£o completos
            - Confirme se campos obrigat√≥rios est√£o preenchidos
            - Considere validar na origem (sistema emissor)
            """)
        else:
            st.success(f"""
            ‚úÖ **Qualidade Excelente!**
            
            Completude de **{completude:.1f}%** indica dados bem estruturados e completos.
            """)
        
        # Detalhamento por tipo (se for MISTO)
        if "MISTO" in tipo_nf:
            st.markdown("---")
            st.subheader("üìã Detalhamento por Tipo de Documento")
            
            col1, col2 = st.columns(2)
            
            with col1:
                # Conta registros de NFe (tem campos emit_ ou cfop)
                registros_nfe = df_unificado[df_unificado.apply(lambda row: any(pd.notna(row[col]) for col in df_unificado.columns if 'emit_' in col or col == 'cfop'), axis=1)]
                st.metric("üîµ Registros NF-e", len(registros_nfe) if not registros_nfe.empty else 0)
            
            with col2:
                # Conta registros de NFSe (tem campos prestador_)
                registros_nfse = df_unificado[df_unificado.apply(lambda row: any(pd.notna(row[col]) for col in df_unificado.columns if 'prestador_' in col), axis=1)]
                st.metric("üü£ Registros NFS-e", len(registros_nfse) if not registros_nfse.empty else 0)
    
    st.markdown("---")

    # Exporta√ß√£o de Relat√≥rio
    st.subheader("üì• Exportar Relat√≥rio de Auditoria")
    
    col1, col2 = st.columns(2)
    
    with col1:
     if st.button("üìÑ Gerar Relat√≥rio PDF", use_container_width=True):
        if st.session_state.get("dados_tabulares"):
            try:
                with st.spinner("Gerando relat√≥rio PDF..."):
                    # Prepara dados da sess√£o
                    dados = {
                        "dados_tabulares": st.session_state.get("dados_tabulares", []),
                        "arquivos_carregados": st.session_state.get("arquivos_carregados", set()),
                        "past": st.session_state.get("past", []),
                        "pdf_list": st.session_state.get("pdf_list", []),
                    }
                    
                    # Gera PDF usando o m√≥dulo externo
                    pdf_bytes = gerar_relatorio_pdf(dados)
                    
                    # Bot√£o de download
                    st.download_button(
                        label="‚¨áÔ∏è Download Relat√≥rio PDF",
                        data=pdf_bytes,
                        file_name=f"relatorio_auditoria_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                    
                    st.success("‚úÖ Relat√≥rio PDF gerado com sucesso!")
                
            except Exception as e:
                st.error(f"‚ùå Erro ao gerar PDF: {str(e)}")
                logger.error(f"Erro ao gerar PDF: {e}")
        else:
            st.warning("‚ö†Ô∏è Nenhum dado dispon√≠vel para gerar relat√≥rio")

    
    with col2:
        if st.button("üìä Exportar para CSV", use_container_width=True):
            if st.session_state.get("dados_tabulares"):
                df_export = pd.concat(st.session_state["dados_tabulares"], ignore_index=True)
                csv = df_export.to_csv(index=False)
                st.download_button(
                    label="‚¨áÔ∏è Download CSV",
                    data=csv,
                    file_name=f"auditoria_fiscal_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv"
                )
            else:
                st.warning("‚ö†Ô∏è Nenhum dado dispon√≠vel para exportar")

# ABA 3: VISUALIZA√á√ïES INTELIGENTES
with abas[3]:
    df_unificado = st.session_state.get("df_csv_unificado")
    
    if df_unificado is not None and not df_unificado.empty:
        
        # üí° ADICIONE ESTE BLOCO AQUI (helper para usu√°rios)
        with st.expander("üí° Como fazer perguntas eficazes", expanded=False):
            st.markdown("""
            ### üìù Dicas para Perguntas Claras:
            
            **Para gr√°fico de PIZZA:**
            - ‚ùå "Mostre em pizza a distribui√ß√£o por CFOP"
            - ‚úÖ "Distribui√ß√£o por CFOP" (e depois escolha pizza)
            
            **Para gr√°fico de BARRAS:**
            - ‚úÖ "Mostre a distribui√ß√£o por CFOP"
            - ‚úÖ "Qual UF tem mais notas?"
            
            **Para gr√°fico de LINHA:**
            - ‚úÖ "Evolu√ß√£o mensal de notas fiscais"
            - ‚úÖ "Tend√™ncia ao longo do tempo"
            
            #### üéØ Colunas Dispon√≠veis:
            """)
            
            # Mostra as colunas dispon√≠veis
            cols_formatadas = ", ".join([f"`{col}`" for col in df_unificado.columns[:10]])
            st.markdown(f"**Principais:** {cols_formatadas}...")
        
        st.markdown("---")
        
        # Seu c√≥digo existente
        exibir_visualizacao(df_unificado)
    else:
        st.warning("üì≠ Nenhum dado carregado. Fa√ßa upload na aba 'Dados & Chat'.")



# ABA 4: PAINEL
with abas[4]:
    st.subheader("üìä Painel Fiscal Inteligente")

    opcao = st.radio(
        "Escolha o que deseja visualizar:",
        [
            "üìä M√©trica CFOP",
            "üìÅ M√©trica SPED",
            "üß† Insights Inteligentes",
            "üîß Sugest√µes de Corre√ß√£o",
            "üì• Exportar Alertas"
        ],
        key="painel_opcao_visualizacao"  # üîß Chave √∫nica e descritiva
    )

    df_unificado = st.session_state.get("df_csv_unificado")

    if df_unificado is not None and not df_unificado.empty:
        if opcao == "üìä M√©trica CFOP":
            from painel_inteligente import gerar_metrica_cfop
            metrica_cfop = gerar_metrica_cfop(df_unificado)
            if metrica_cfop.empty:
                st.warning("‚ö†Ô∏è Nenhuma coluna de CFOP encontrada.")
            else:
                st.dataframe(metrica_cfop)
                import plotly.express as px
                fig = px.bar(metrica_cfop, x='CFOP', y='Quantidade', title='Distribui√ß√£o de CFOPs')
                st.plotly_chart(fig)

        elif opcao == "üìÅ M√©trica SPED":
            from painel_inteligente import gerar_metrica_sped
            metrica_sped = gerar_metrica_sped(df_unificado)
            if metrica_sped.empty:
                st.warning("‚ö†Ô∏è Nenhuma coluna relacionada ao SPED foi encontrada.")
            else:
                st.dataframe(metrica_sped)

        elif opcao == "üß† Insights Inteligentes":
            from painel_inteligente import analise_inteligente
            insights = analise_inteligente(df_unificado)
            for item in insights:
                st.write(item)

        elif opcao == "üîß Sugest√µes de Corre√ß√£o":
            from painel_inteligente import sugerir_correcao
            sugestoes = sugerir_correcao(df_unificado)
            for item in sugestoes:
                st.write(item)

        elif opcao == "üì• Exportar Alertas":
            from painel_inteligente import gerar_alertas
            alertas = gerar_alertas(df_unificado)
            if alertas:
                for alerta in alertas:
                    st.warning(alerta)

                import io
                from docx import Document
                doc = Document()
                doc.add_heading("Alertas Fiscais", level=1)
                for alerta in alertas:
                    doc.add_paragraph(alerta)

                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.download_button("üì• Baixar alertas em Word", buffer, file_name="alertas_fiscais.docx")
            else:
                st.success("‚úÖ Nenhum alerta fiscal encontrado.")
    else:
        st.info("üì≠ Nenhum dado fiscal carregado. Fa√ßa upload na aba 'Dados & Chat'.")

# Rodap√©
try:
    exibir_rodape()
except:
    st.markdown("---")
    st.caption("ChatFiscal 2025 | Com Mem√≥ria Inteligente")